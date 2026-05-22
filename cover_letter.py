#!/usr/bin/env python3
"""
============================================================
  cover_letter.py  —  Cover letter generator
============================================================

HOW TO RUN:
    1. Paste the full job posting into:    job-apps/job_description.txt
    2. Save your professional bio in:      job-apps/bio.txt
    3. Open a terminal in the job-apps folder and run:
           python cover_letter.py

AI MODE (strongly recommended — much better output):
    Set your Anthropic API key before running:

        Windows PowerShell:  $env:ANTHROPIC_API_KEY = "sk-ant-..."
        Windows CMD:         set ANTHROPIC_API_KEY=sk-ant-...
        Mac / Linux:         export ANTHROPIC_API_KEY=sk-ant-...

    Get a free key at: https://console.anthropic.com

    Without a key: a structured template is filled using
    keywords and bullet points extracted from the JD and bio.
    All [EDIT: ...] markers must be replaced before sending.

OUTPUT:  output/cover_letter_[Company]_[YYYY-MM-DD].txt

WHAT TO PUT IN bio.txt:
    2–4 paragraphs about yourself — rough notes are fine:
      • Total years of experience and professional background
      • Key skills and areas of specialization
      • Notable achievements, projects, or clients
      • What motivates you / what kind of role you're seeking

REQUIREMENTS:  Python 3.8+, no third-party packages needed.
============================================================
"""

import os
import sys
import re
import json
import datetime
import urllib.request
import urllib.error
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

# ── File paths ────────────────────────────────────────────────────────────────
SCRIPT_DIR    = Path(__file__).parent.resolve()
OUTPUT_DIR    = SCRIPT_DIR / "output"
JD_FILE       = SCRIPT_DIR / "job_description.txt"
BIO_FILE      = SCRIPT_DIR / "bio.txt"
MASTER_JSON   = SCRIPT_DIR / "master_resume.json"
RESUMES_DIR        = SCRIPT_DIR / "resumes"
COVER_LETTERS_DIR  = SCRIPT_DIR / "Cover Letters"


# ── Helpers ───────────────────────────────────────────────────────────────────

def guess_company(jd_text: str) -> str:
    """Attempt to auto-detect company name from the job description."""
    patterns = [
        r"(?:Company|Employer|Organization):\s*([^\n]{2,60})",
        r"^([A-Z][a-zA-Z0-9&\s]{2,40}?)\s+is\s+(?:a|an|the)\b",
        r"(?:at|join|about)\s+([A-Z][a-zA-Z0-9&\s]{1,30}?)(?:\s+is|\s+are|\.|,)",
    ]
    for pat in patterns:
        m = re.search(pat, jd_text, re.MULTILINE)
        if m:
            name = m.group(1).strip().rstrip(".,;")
            if 2 < len(name) < 60:
                return name
    return "the company"


def guess_role(jd_text: str) -> str:
    """Attempt to auto-detect the job title from the job description."""
    patterns = [
        r"(?:Job Title|Position|Role|Title):\s*([^\n]{2,80})",
        r"(?:We are (?:looking for|hiring|seeking)|We need)\s+(?:a|an)\s+([^\n.]{5,80})",
        r"^([A-Z][a-zA-Z\s/–\-]{5,60})\s*$",
    ]
    for pat in patterns:
        m = re.search(pat, jd_text, re.MULTILINE | re.IGNORECASE)
        if m:
            role = m.group(1).strip().rstrip(".,;")
            if 4 < len(role) < 80:
                return role
    return "this position"


def safe_slug(name: str) -> str:
    """Convert a name to a safe filename segment."""
    return re.sub(r"[^\w]", "_", name).strip("_")[:40]


def extract_top_requirements(jd_text: str, n: int = 5) -> list:
    """Pull the first n bullet-point requirements from the JD."""
    bullets = re.findall(r'(?:^|\n)\s*[•\-\*]\s*(.+)', jd_text)
    if not bullets:
        # Fall back to sentences containing "must" or "required"
        bullets = re.findall(
            r'[^.]*\b(?:must|required)\b[^.]*\.', jd_text, re.IGNORECASE
        )
    return [b.strip() for b in bullets[:n] if len(b.strip()) > 15]


def first_sentences(text: str, n: int = 2) -> list:
    """Return the first n non-empty sentences from a block of text."""
    sentences = [s.strip() for s in re.split(r'(?<=[.!?])\s+', text) if len(s.strip()) > 20]
    return sentences[:n]


# ── API-powered generation ────────────────────────────────────────────────────

def call_api(prompt: str, api_key: str) -> str:
    """POST to the Anthropic Messages API using only urllib (no pip installs)."""
    payload = json.dumps({
        "model": "claude-sonnet-4-6",
        "max_tokens": 2000,
        "messages": [{"role": "user", "content": prompt}],
    }).encode("utf-8")

    req = urllib.request.Request(
        "https://api.anthropic.com/v1/messages",
        data=payload,
        headers={
            "x-api-key": api_key,
            "anthropic-version": "2023-06-01",
            "content-type": "application/json",
        },
    )
    try:
        with urllib.request.urlopen(req, timeout=90) as resp:
            body = json.loads(resp.read().decode("utf-8"))
            return body["content"][0]["text"]
    except urllib.error.HTTPError as e:
        raise RuntimeError(f"API error {e.code}: {e.read().decode('utf-8')[:400]}")


def generate_with_api(bio_text: str, jd_text: str, company: str, role: str, api_key: str) -> str:
    """Ask Claude to write a tailored cover letter body (no letterhead — added separately in DOCX)."""
    today = datetime.date.today().strftime("%B %d, %Y")
    prompt = (
        "Write a professional cover letter for the job below.\n\n"
        "RULES:\n"
        "1. Address it to 'Hiring Team' — we don't know the hiring manager's name.\n"
        "2. Opening paragraph: express genuine enthusiasm for THIS specific role and company.\n"
        "   Reference something specific about the company — their product, mission, growth, or culture.\n"
        "   No generic openers like 'I am writing to apply for...'.\n"
        "3. Body (2 paragraphs): draw direct, specific connections between the candidate's\n"
        "   background and the role's key requirements. Mirror the JD's own language and keywords.\n"
        "   Reference actual requirements from the JD by name.\n"
        "4. Closing: confident and warm. Invite next steps. No desperate or sycophantic language.\n"
        "5. Length: 3–4 paragraphs, under 400 words total.\n"
        "6. Format: plain text. Start with the date on the first line, then a blank line, then the\n"
        "   company/team name, then 'Dear Hiring Team,', then body paragraphs separated by blank lines,\n"
        "   then 'Sincerely,' on its own line, then a blank line, then 'Kiara Earl' on the last line.\n"
        "   Do NOT include a letterhead or contact info block — that is added separately.\n"
        "7. Output ONLY the letter body as described — no commentary, no markdown fencing.\n"
        "8. Never use an em dash (—) inside parentheses. Rewrite using a comma, colon, or separate sentence.\n\n"
        f"DATE: {today}\n"
        f"COMPANY: {company}\n"
        f"ROLE: {role}\n\n"
        f"CANDIDATE BIO:\n{bio_text}\n\n"
        f"JOB DESCRIPTION:\n{jd_text[:3000]}\n\n"
        "Cover letter:"
    )
    return call_api(prompt, api_key)


# ── Fallback: template-based generation ──────────────────────────────────────

def generate_template(bio_text: str, jd_text: str, company: str, role: str) -> str:
    """Fill a structured template when no API key is available."""
    today = datetime.date.today().strftime("%B %d, %Y")
    sig_name, sig_contact = _load_contact_info()
    requirements = extract_top_requirements(jd_text)
    req1 = requirements[0] if requirements else "the key responsibilities in the posting"
    req2 = requirements[1] if len(requirements) > 1 else "the team's goals"

    bio_sentences = first_sentences(bio_text, n=3)
    strength1 = bio_sentences[0] if bio_sentences else "[Your key professional strength]"
    strength2 = bio_sentences[1] if len(bio_sentences) > 1 else "[Another relevant strength]"

    letter = f"""{today}

Hiring Team
{company}

Dear Hiring Team,

[EDIT: Replace this sentence with a specific, enthusiastic opener. Why THIS company? Reference
their product, mission, recent news, or something you genuinely admire. Avoid "I am writing to
apply..."] I am excited to apply for the {role} position at {company}.

{strength1}. This background positions me well to contribute to {req1.lower()}. [EDIT: Add a
concrete example or measurable achievement that proves this.] Beyond that, {strength2.lower()},
which I believe maps directly to {req2.lower()}. [EDIT: Expand with a specific project, result,
or metric from your experience.]

[EDIT: Write a third paragraph connecting one more aspect of your experience to the role. Good
candidates reference a specific tool, methodology, or accomplishment from their resume that
matches something in the job description. Be specific — generic claims don't stand out.]

Thank you for taking the time to review my application. I would welcome the opportunity to discuss
how my background aligns with {company}'s needs and would be glad to connect at your convenience.

Sincerely,

{sig_name}
{sig_contact}"""

    return letter


# ── DOCX builder ─────────────────────────────────────────────────────────────

def _load_contact_info() -> tuple:
    """Return (name, contact_line) from master_resume.json, with hardcoded fallback."""
    try:
        if MASTER_JSON.exists():
            data = json.loads(MASTER_JSON.read_text(encoding="utf-8"))
            name = data.get("name", "")
            parts = [
                data.get("location", ""),
                data.get("phone", ""),
                data.get("email", ""),
                data.get("linkedin", ""),
            ]
            contact = " | ".join(p for p in parts if p)
            if name and contact:
                return name, contact
    except Exception:
        pass
    return (
        "Kiara Earl",
        "Houston, TX | 346-777-3451 | kimearls24@outlook.com | linkedin.com/in/kiaraearl",
    )


def build_cover_letter_docx(letter_body: str, company: str, out_dir: Path) -> Path:
    """
    Build a DOCX cover letter matching the reference format:
      - Name bold at top, contact info below (8pt space after each line)
      - Each letter paragraph on its own line with 8pt spacing after
      - Calibri 11pt throughout; name is 13pt bold
    """
    if not _DOCX_AVAILABLE:
        raise RuntimeError("python-docx is not installed. Run: pip install python-docx")

    _FONT    = "Calibri"
    _BODY_PT = 11.0
    _NAME_PT = 13.0
    _SPC_AFT = 8.0   # space after each paragraph — matches reference (160 twips ÷ 20)

    name, contact_line = _load_contact_info()

    doc = Document()

    for sec in doc.sections:
        sec.top_margin    = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin   = Inches(1)
        sec.right_margin  = Inches(1)

    nstyle = doc.styles["Normal"]
    nstyle.font.name = _FONT
    nstyle.font.size = Pt(_BODY_PT)

    def _para(text, bold=False, size=_BODY_PT, space_after=_SPC_AFT):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(space_after)
        r = p.add_run(text)
        r.font.name  = _FONT
        r.font.size  = Pt(size)
        r.font.bold  = bold
        return p

    # Letterhead
    _para(name, bold=True, size=_NAME_PT)
    _para(contact_line)

    # Letter body — split on blank lines to get paragraphs
    blocks = [b.strip() for b in re.split(r"\n{2,}", letter_body.strip()) if b.strip()]
    for block in blocks:
        # Collapse internal newlines (wrap from API)
        text = " ".join(line.strip() for line in block.splitlines() if line.strip())
        _para(text)

    out_dir.mkdir(exist_ok=True)
    filename = f"KiaraEarl_{safe_slug(company)}_CoverLetter.docx"
    out_path = out_dir / filename
    doc.save(str(out_path))
    return out_path


def generate_cover_letter_docx(jd_text: str, company: str, role: str,
                                api_key: str, out_dir: Path = None) -> dict:
    """
    Full pipeline: call Claude → build DOCX.

    Returns:
      {"filename": str, "error": str or None}
    """
    if out_dir is None:
        out_dir = COVER_LETTERS_DIR

    bio_path = BIO_FILE
    bio_text = ""
    try:
        bio_text = bio_path.read_text(encoding="utf-8").strip() if bio_path.exists() else ""
    except Exception:
        pass

    used_template = False
    try:
        letter_body = generate_with_api(bio_text or "See resume.", jd_text, company, role, api_key)
    except Exception as exc:
        print(f"[cover_letter] API call failed ({type(exc).__name__}): {exc}")
        letter_body = generate_template(bio_text or "", jd_text, company, role)
        used_template = True

    try:
        out_path = build_cover_letter_docx(letter_body, company, out_dir)
        result = {"filename": out_path.name, "error": None}
        if used_template:
            result["warning"] = "template_used"
        return result
    except Exception as exc:
        return {"filename": None, "error": str(exc)}


# ── Entry point ───────────────────────────────────────────────────────────────

def main():
    print("\n=== Cover Letter Generator ===\n")

    # Validate input files
    for path, name in [(JD_FILE, "job_description.txt"), (BIO_FILE, "bio.txt")]:
        if not path.exists():
            sys.exit(f"ERROR: {name} not found.\nExpected: {path}")
        if not path.read_text(encoding="utf-8").strip():
            sys.exit(f"ERROR: {name} is empty — add your content and try again.")

    jd_text  = JD_FILE.read_text(encoding="utf-8").strip()
    bio_text = BIO_FILE.read_text(encoding="utf-8").strip()

    # Company and role
    auto_company = guess_company(jd_text)
    entered_co   = input(f"Company name [{auto_company}]: ").strip()
    company      = entered_co or auto_company

    auto_role   = guess_role(jd_text)
    entered_rol = input(f"Role / job title [{auto_role}]: ").strip()
    role        = entered_rol or auto_role

    # Generate
    api_key = os.environ.get("ANTHROPIC_API_KEY", "").strip()
    if api_key:
        print("\nAPI key found — generating with Claude...")
        try:
            letter = generate_with_api(bio_text, jd_text, company, role, api_key)
            mode = "AI-generated"
        except Exception as exc:
            print(f"Warning: API call failed ({exc})\nFalling back to template...")
            letter = generate_template(bio_text, jd_text, company, role)
            mode = "template (API failed)"
    else:
        print("\nNo ANTHROPIC_API_KEY found — using template mode.")
        print("Tip: set the key for AI-generated letters (see script header).")
        letter = generate_template(bio_text, jd_text, company, role)
        mode = "template"

    # Write output
    OUTPUT_DIR.mkdir(exist_ok=True)
    out_path = OUTPUT_DIR / f"cover_letter_{safe_slug(company)}_{datetime.date.today()}.txt"
    out_path.write_text(letter, encoding="utf-8")

    print(f"\n[{mode}] Saved: {out_path}")
    if "template" in mode:
        print("Edit all [EDIT: ...] markers before submitting.")


if __name__ == "__main__":
    main()
