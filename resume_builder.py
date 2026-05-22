#!/usr/bin/env python3
"""
resume_builder.py — ATS-optimized DOCX resume generator

Workflow:
  1. Parse the job description for keywords
  2. Cross-reference master_resume.json to calculate ATS score
  3. If score >= 80, call Claude to tailor content, then build DOCX
  4. Return score, matched/missing keywords, and output file path
"""

import json
import re
import collections
import datetime
import urllib.request
import urllib.error
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import lxml.etree as etree
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

SCRIPT_DIR   = Path(__file__).parent.resolve()
MASTER_JSON  = SCRIPT_DIR / "master_resume.json"
RESUMES_DIR  = SCRIPT_DIR / "resumes"

# Employment dates — hardcoded permanently; never pulled from master_resume.json
HARDCODED_DATES = {
    "Progressive Insurance": "Apr 2025 - Present",
    "World Travel Holdings": "Apr 2022 - Apr 2024",
    "Bluegreen Vacations":   "Sept 2021 - Apr 2022",
}


STOPWORDS = {
    "a","an","the","and","or","but","in","on","at","to","for","of","with",
    "by","from","up","about","into","is","are","was","were","be","been",
    "being","have","has","had","do","does","did","will","would","could",
    "should","may","might","must","shall","can","we","you","they","he",
    "she","it","this","that","these","those","i","my","our","your","their",
    "its","who","which","what","when","where","why","how","all","both",
    "each","few","more","most","other","some","such","no","not","only",
    "same","so","than","too","very","just","now","as","if","while","also",
    "well","new","key","please","us","position","role","team","years","year",
    "experience","ability","skills","knowledge","work","working","opportunity",
    "company","looking","seeking","required","requirements","preferred",
    "qualifications","responsibilities","duties","include","including",
    "following","related","relevant","demonstrated","proven","able","plus",
    "etc","per","via","use","using","used","get","make","help","ensure",
    "provide","support","strong","excellent","good","great","high","person",
    "candidates","candidate","applicant","applicants","background","degree",
    "responsible","day","days","including","monday","friday","full","time",
    "part","based","join","across","within","multiple","various","various",
    "first","second","third","fourth","fifth","like","make","take","bring",
    "ensure","maintain","develop","build","create","manage","lead","drive",
    "identify","implement","improve","review","report","track","monitor",
}


def load_master() -> dict:
    """Load and return master_resume.json as a dict."""
    if not MASTER_JSON.exists():
        return {}
    return json.loads(MASTER_JSON.read_text(encoding="utf-8"))


def extract_jd_keywords(jd_text: str, top_n: int = 50,
                         company: str = None) -> list:
    """
    Extract the most important keywords from a job description.
    Returns a ranked list of keyword strings.
    """
    # Build a per-call exclusion set: stopwords + company name tokens
    exclusions = set(STOPWORDS)
    if company:
        for tok in re.findall(r"[a-zA-Z]+", company):
            if len(tok) > 2:
                exclusions.add(tok.lower())

    words = re.findall(r"\b[a-zA-Z][a-zA-Z0-9+#.]*\b", jd_text)
    freq: dict = collections.defaultdict(float)

    for w in words:
        wl = w.lower()
        if wl not in exclusions and len(wl) > 2:
            freq[wl] += 1.0

    # Boost keywords in requirements sections
    req_match = re.search(
        r'(?:required|qualifications?|must.have|requirements?|preferred)[\s\S]*?(?=\n\n|\Z)',
        jd_text, re.IGNORECASE,
    )
    if req_match:
        for w in re.findall(r"\b[a-zA-Z][a-zA-Z0-9+#.]*\b", req_match.group()):
            wl = w.lower()
            if wl not in STOPWORDS and len(wl) > 2:
                freq[wl] += 2.0

    # Boost common tool/skill terms
    tech_pattern = re.compile(
        r'\b(?:python|javascript|java|sql|excel|salesforce|quickbooks|netsuite|sap|'
        r'oracle|tableau|powerbi|power bi|microsoft|office|word|outlook|teams|zoom|'
        r'slack|jira|confluence|hubspot|zendesk|servicenow|workday|adp|paychex|'
        r'healthcare|hipaa|phi|ehr|emr|cpt|icd|billing|coding|compliance|'
        r'data entry|accounts payable|accounts receivable|bookkeeping|payroll|'
        r'scheduling|documentation|customer service|customer support|crm|'
        r'virtual assistant|project management|process improvement|'
        r'analytical|communication|organizational|attention|detail)\b',
        re.IGNORECASE,
    )
    for m in tech_pattern.findall(jd_text):
        wl = m.lower()
        if wl not in STOPWORDS:
            freq[wl] = freq.get(wl, 0) + 2.0

    # Include two-word phrases that appear more than once
    clean = [w.lower() for w in words if w.lower() not in STOPWORDS and len(w) > 2]
    bigrams = collections.Counter(
        f"{clean[i]} {clean[i+1]}" for i in range(len(clean) - 1)
    )
    for bg, cnt in bigrams.items():
        if cnt >= 2:
            freq[bg] = freq.get(bg, 0) + cnt * 1.5

    return [kw for kw, _ in sorted(freq.items(), key=lambda x: -x[1])[:top_n]]


def ats_score(jd_text: str, master: dict, tailored_text: str = None,
              company: str = None) -> dict:
    """
    Calculate ATS keyword match score.
    If tailored_text is provided, score against that instead of rebuilding from master.

    Returns:
      score      int (0-100)
      matched    list of keyword strings found in resume
      missing    list of keyword strings not found
      total      int total keywords evaluated
    """
    keywords = extract_jd_keywords(jd_text, top_n=40, company=company)

    if tailored_text is not None:
        resume_blob = tailored_text.lower()
    else:
        # Build a searchable blob from master resume fields
        parts = []
        parts.extend(master.get("skills", []))
        parts.extend(master.get("track_2_skills", []))
        for exp in master.get("experience", []):
            parts.append(exp.get("title", ""))
            parts.append(exp.get("company", ""))
            parts.extend(exp.get("bullets", []))
        for cert in master.get("certifications", []):
            parts.append(cert)
        for edu in master.get("education", []):
            parts.append(edu.get("degree", ""))
            parts.append(edu.get("school", ""))
        resume_blob = " ".join(parts).lower()

    matched = []
    missing = []
    for kw in keywords:
        if kw in resume_blob:
            matched.append(kw)
        else:
            missing.append(kw)

    total = len(keywords)
    score = round(len(matched) / total * 100) if total else 0

    return {
        "score":   score,
        "matched": matched,
        "missing": missing,
        "total":   total,
    }


def _call_api(prompt: str, api_key: str, model: str = "claude-sonnet-4-6") -> str:
    """POST to Anthropic Messages API using urllib."""
    payload = json.dumps({
        "model": model,
        "max_tokens": 4000,
        "messages": [{"role": "user", "content": prompt}],
    }).encode("utf-8")

    req = urllib.request.Request(
        "https://api.anthropic.com/v1/messages",
        data=payload,
        headers={
            "x-api-key": api_key,
            "anthropic-version": "2023-06-01",
            "content-type": "application/json",
        },
    )
    with urllib.request.urlopen(req, timeout=120) as resp:
        body = json.loads(resp.read().decode("utf-8"))
        return body["content"][0]["text"]


def tailor_docx_content(jd_text: str, master: dict, company: str,
                         job_title: str, api_key: str) -> dict:
    """
    Ask Claude to produce a JSON payload with tailored resume content.

    Returns dict with keys: summary, competencies (list), bullets_by_role (dict),
    selected_projects (list of project ids).
    """
    skills_str = ", ".join(master.get("skills", []) + master.get("track_2_skills", []))
    exp_list = []
    for exp in master.get("experience", []):
        exp_list.append(f"{exp.get('title','')} at {exp.get('company','')} ({exp.get('dates','')})")
    certs_str = "; ".join(master.get("certifications", []))
    projects_str = "\n".join(
        f'- id: "{p["id"]}" | title: "{p["title"]}" | tags: {", ".join(p.get("tags", []))}'
        for p in master.get("projects", [])
    )

    prompt = (
        "You are an expert ATS resume writer. Your goal is 80%+ keyword match against the job description.\n\n"
        "OUTPUT FORMAT — respond with ONLY valid JSON, no markdown fencing:\n"
        "{\n"
        '  "summary": "<3-5 sentences, third-person implied (no I), uses the EXACT job title from the posting, '
        'mirrors 3-5 keywords verbatim from the JD>",\n'
        '  "competencies": ["<keyword exactly as it appears in the JD>", ...],\n'
        '  "bullets_by_role": {\n'
        '    "Progressive Insurance": ["<bullet 1>","<bullet 2>","<bullet 3>","<bullet 4>","<bullet 5>"],\n'
        '    "World Travel Holdings": ["<bullet 1>","<bullet 2>","<bullet 3>","<bullet 4>"],\n'
        '    "Bluegreen Vacations": ["<bullet 1>","<bullet 2>"]\n'
        '  },\n'
        '  "include_bluegreen": <true ONLY if Bluegreen Vacations experience (inbound phone support, '
        'conflict resolution, billing/scheduling) directly strengthens the application vs. omitting it; '
        'false if Progressive + World Travel already cover all competitive keywords>,\n'
        '  "selected_projects": [\n'
        '    {"id": "<project id>", "description": "<rewritten description>", "tech": "<rewritten tech line>"}\n'
        '  ]\n'
        "  Pick 0-3 projects from AVAILABLE PROJECTS. Empty array [] for pure sales/admin/billing/VA roles.\n"
        "  For IT/helpdesk/cybersecurity/tech-adjacent roles pick whichever projects best match JD keywords.\n"
        "}\n\n"
        "RULES:\n"
        "1. Never fabricate experience, employers, or credentials.\n"
        "2. competencies: 10-14 keywords. Pull them VERBATIM from the JD — exact spelling and casing.\n"
        "3. bullets_by_role: Progressive Insurance gets EXACTLY 5 bullets, World Travel Holdings gets "
        "EXACTLY 4 bullets, Bluegreen Vacations gets EXACTLY 2 bullets. Each bullet is ONE sentence. "
        "Use the JD's own words where possible. "
        "Do not start two consecutive bullets with the same action verb. "
        "At least one bullet per role must include a number or percentage.\n"
        "4. summary: no 'I' statements. Must include the exact job title and at least 3 terms lifted "
        "directly from the Requirements or Qualifications section of the JD.\n"
        "5. ATS target: 80%+ of the top 30 keywords from the JD must appear somewhere in the resume output.\n"
        "6. The final resume must fit on ONE page. Keep bullets tight — one strong sentence each, "
        "no padding, no filler phrases.\n"
        "7. NEVER include any language suggesting this resume or cover letter was written, generated, "
        "or assisted by AI. No phrases like 'AI-generated', 'AI-assisted', 'created with AI', etc.\n"
        "8. Never use an em dash (—) inside parentheses. Rewrite using a comma, colon, or separate sentence.\n"
        "9. Authentic role representation: all three roles had a genuine sales component — preserve that truth "
        "in the bullets so the resume matches verifiable work history. Each role must include AT LEAST 1-2 bullets "
        "that reflect the real sales or revenue-generating nature of the position: "
        "Progressive Insurance = multi-product insurance sales, cross-selling, and quota/target attainment; "
        "World Travel Holdings = travel package sales alongside technical support; "
        "Bluegreen Vacations = vacation ownership sales, closing, and client consultation. "
        "The remaining bullets can pivot toward tech-adjacent or transferable skills that match the JD, "
        "but the sales bullets anchor the role as real and verifiable.\n"
        "10. selected_projects rewriting rules: reword each selected project's description and tech line "
        "to mirror the JD's language and emphasize the most relevant aspects. "
        "You MAY reorder emphasis, swap synonyms, and lead with JD-matching keywords. "
        "You MUST NOT add tools, technologies, certifications, or outcomes that were not in the original. "
        "You MUST NOT fabricate metrics or numbers. Core facts (what was done, what tools were used) must remain accurate.\n\n"
        f"COMPANY: {company}\n"
        f"JOB TITLE: {job_title}\n\n"
        f"JOB DESCRIPTION:\n{jd_text[:3000]}\n\n"
        f"CANDIDATE SKILLS: {skills_str}\n"
        f"CANDIDATE EXPERIENCE: {'; '.join(exp_list)}\n"
        f"CANDIDATE CERTS: {certs_str}\n"
        f"AVAILABLE PROJECTS (select by id):\n{projects_str}\n"
    )

    raw = _call_api(prompt, api_key)
    raw = re.sub(r"^```(?:json)?\s*|\s*```$", "", raw, flags=re.DOTALL).strip()
    return json.loads(raw)


_ABBREVS = {
    "AI", "API", "ATS", "AWS", "B2B", "B2C", "CRM", "CSS", "CX",
    "DNS", "ERP", "GCP", "HR", "HRIS", "HTML", "HTTP", "HTTPS",
    "IT", "ITSM", "KPI", "KYC", "OKR", "PCI", "PPC", "QA", "ROI",
    "SDR", "SEO", "SLA", "SMB", "SME", "SOC", "SOP", "SQL", "SSO",
    "TCP", "UI", "URL", "UX", "VPN", "WFH",
}
# Mixed-case abbreviations that aren't all-caps
_MIXED_ABBREVS = {"voip": "VoIP", "saas": "SaaS", "paas": "PaaS", "iaas": "IaaS"}


def _smart_comp_case(phrase: str) -> str:
    """Title-case a competency phrase but restore known abbreviations."""
    words = []
    for w in phrase.split():
        upper = w.upper()
        if upper in _ABBREVS:
            words.append(upper)
        elif w.lower() in _MIXED_ABBREVS:
            words.append(_MIXED_ABBREVS[w.lower()])
        else:
            words.append(w.title())
    return " ".join(words)


def build_docx(content: dict, master: dict, company: str,
               job_title: str, ats_data: dict, out_dir: Path) -> Path:
    """
    Build an ATS-safe DOCX using the exact template spec.

    Layout:
      Contact Block (name + one pipe-separated contact line)
      Professional Summary
      Core Competencies  (one pipe-separated line)
      Professional Experience  (tab layout: title | company [right-tab] italic date)
      Education & Certifications  (hardcoded fixed entries)
      Projects  (conditional — tech/IT roles only)

    ATS rules: Calibri, single column, no tables, no headers/footers,
    no images, no horizontal lines, 1-inch margins.
    """
    if not _DOCX_AVAILABLE:
        raise RuntimeError("python-docx is not installed. Run: pip install python-docx")

    _FONT      = "Calibri"
    _BODY_PT   = 10.5
    _HEADER_PT = 11.0
    _NAME_PT   = 18.0
    _TW        = 7.0   # text width in inches (8.5" page, 0.75" margins each side)

    doc = Document()

    # ── 0.75-inch margins (tighter for 1-page fit) ────────────────────────────
    for sec in doc.sections:
        sec.top_margin    = Inches(0.75)
        sec.bottom_margin = Inches(0.75)
        sec.left_margin   = Inches(0.75)
        sec.right_margin  = Inches(0.75)
        sec.header_distance = Inches(0)
        sec.footer_distance = Inches(0)

    # Base Normal style — Calibri 10.5pt
    nstyle = doc.styles["Normal"]
    nstyle.font.name = _FONT
    nstyle.font.size = Pt(_BODY_PT)

    # ── Helpers ───────────────────────────────────────────────────────────────

    def _new_para(space_before=0, space_after=2, left_indent=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        if left_indent is not None:
            p.paragraph_format.left_indent = Inches(left_indent)
        return p

    def _run(para, text, bold=False, italic=False, size=_BODY_PT):
        r = para.add_run(text)
        r.font.name   = _FONT
        r.font.size   = Pt(size)
        r.font.bold   = bold
        r.font.italic = italic
        return r

    def _right_tab(para):
        """Add a right-aligned tab stop at the right text boundary."""
        pPr = para._p.get_or_add_pPr()
        tabs_el = pPr.find(qn("w:tabs"))
        if tabs_el is None:
            tabs_el = OxmlElement("w:tabs")
            pPr.append(tabs_el)
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "right")
        tab.set(qn("w:pos"), str(int(_TW * 1440)))
        tabs_el.append(tab)

    def _add_bottom_border(para):
        """Add a thin bottom border under a paragraph (used on section headers)."""
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"),   "single")
        bottom.set(qn("w:sz"),    "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "000000")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _section_header(label: str):
        """Bold, 11pt, with a bottom border rule — matches reference docx."""
        p = _new_para(space_before=8, space_after=2)
        _run(p, label, bold=True, size=_HEADER_PT)
        _add_bottom_border(p)

    def _body_line(text: str, bold=False, space_after=2):
        p = _new_para(space_after=space_after)
        _run(p, text, bold=bold)
        return p

    def _tab_role_line(title: str, rest: str, date_str: str,
                       space_before=6, space_after=2):
        """
        Formats: [bold title] | [rest]  [right-tab][italic date]
        Used for both experience and education entries.
        """
        p = _new_para(space_before=space_before, space_after=space_after)
        _right_tab(p)
        _run(p, title, bold=True)
        if rest:
            _run(p, f" | {rest}")
        if date_str:
            _run(p, "\t")                       # push date to right margin
            _run(p, date_str, italic=True)
        return p

    def _bullet(text: str):
        """List Bullet style with tight spacing — matches reference docx."""
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        r = p.add_run(text)
        r.font.name = _FONT
        r.font.size = Pt(_BODY_PT)
        return p

    # ── Contact Block (centered) ──────────────────────────────────────────────
    name_p = _new_para(space_before=0, space_after=2)
    name_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(name_p, master.get("name", ""), bold=True, size=16.0)

    role_p = _new_para(space_before=0, space_after=2)
    role_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(role_p, job_title, size=13.0)

    contact_parts = [
        master.get("location", ""),
        master.get("phone", ""),
        master.get("email", ""),
        master.get("linkedin", ""),
    ]
    contact_line = " | ".join(p for p in contact_parts if p)
    contact_p = _new_para(space_after=4)
    contact_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(contact_p, contact_line)

    # ── Professional Summary ──────────────────────────────────────────────────
    _section_header("Professional Summary")
    summary = content.get("summary", "")
    if summary:
        _body_line(summary, space_after=4)

    # ── Core Competencies (3 per row) ─────────────────────────────────────────
    _section_header("Core Competencies")
    comps = [_smart_comp_case(c) for c in content.get("competencies", [])]
    if comps:
        rows = [comps[i:i+3] for i in range(0, len(comps), 3)]
        for i, row in enumerate(rows):
            after = 4 if i == len(rows) - 1 else 1
            _body_line("   |   ".join(row), space_after=after)

    # ── Professional Experience ───────────────────────────────────────────────
    _section_header("Professional Experience")
    bullets_by_role  = content.get("bullets_by_role", {})
    include_bluegreen = content.get("include_bluegreen", True)

    # Per-role bullet caps: current job gets 5, WTH gets 4, Bluegreen gets 2
    _bullet_caps = {
        "Progressive Insurance":  5,
        "World Travel Holdings":  4,
        "Bluegreen Vacations":    2,
    }

    for exp in master.get("experience", []):
        co    = exp.get("company", "")
        title = exp.get("title", "")
        loc   = exp.get("location", "") or "Remote"
        dates = HARDCODED_DATES.get(co, exp.get("dates", ""))

        if co == "Bluegreen Vacations" and not include_bluegreen:
            continue

        # Line 1: [bold Company | Location]  [right-tab][italic Date]
        _tab_role_line(co, loc, dates, space_before=8, space_after=0)
        # Line 2: [italic Title] — no right tab
        p = _new_para(space_before=0, space_after=2)
        _run(p, title, italic=True)

        cap = _bullet_caps.get(co, 4)
        role_bullets = (
            bullets_by_role.get(co)
            or bullets_by_role.get(title)
            or exp.get("bullets", [])
        )
        for b in role_bullets[:cap]:
            _bullet(b)
        if not role_bullets:
            _bullet("Responsibilities and achievements tailored to the target role.")

    # ── Projects (selective — Claude picks and rewrites 0-3) ─────────────────────
    projects_by_id = {p["id"]: p for p in master.get("projects", [])}
    selected_projects = [p for p in content.get("selected_projects", []) if isinstance(p, dict)]
    if selected_projects:
        _section_header("Projects")
        for proj in selected_projects:
            pid = proj.get("id", "")
            title = projects_by_id.get(pid, {}).get("title", "")
            description = proj.get("description", projects_by_id.get(pid, {}).get("description", ""))
            tech = proj.get("tech", projects_by_id.get(pid, {}).get("tech", ""))
            if not title:
                continue
            proj_p = _new_para(space_before=4, space_after=2)
            _run(proj_p, title, bold=True)
            _body_line(description, space_after=2)
            _body_line(tech, space_after=2)

    # ── Education & Certifications (fixed template entries) ───────────────────
    _section_header("Education & Certifications")

    # Entry 1 — WGU B.S.
    _tab_role_line(
        "B.S. Cybersecurity & Information Assurance",
        "Western Governors University",
        "Expected 2027",
        space_before=4, space_after=1,
    )

    # Entry 2 — Lone Star A.A.
    _tab_role_line(
        "A.A. General Arts",
        "Lone Star College",
        "Graduated 2018",
        space_before=3, space_after=1,
    )

    # Entry 3 — Certs on one pipe-separated line (status embedded, no date tab)
    certs_p = _new_para(space_before=3, space_after=4, left_indent=0.25)
    _run(certs_p,
         "CompTIA A+ - Active  |  CompTIA Network+ - In Progress"
         "  |  Personal Lines Insurance License - Active")

    # ── Save ──────────────────────────────────────────────────────────────────
    out_dir.mkdir(exist_ok=True)

    def _safe(s: str) -> str:
        return re.sub(r"[^\w]", "_", s).strip("_")[:30]

    filename = f"KiaraEarl_{_safe(company)}_Resume.docx"
    out_path = out_dir / filename
    doc.save(str(out_path))
    return out_path


def generate_ats_docx(jd_text: str, company: str, job_title: str,
                       api_key: str, out_dir: Path = None) -> dict:
    """
    Full pipeline: score → tailor → build DOCX.

    Returns:
      {
        "ats_score":   int,
        "matched":     list,
        "missing":     list,
        "total":       int,
        "file":        Path or None,
        "filename":    str or None,
        "below_threshold": bool,
        "error":       str or None,
      }
    """
    if out_dir is None:
        out_dir = RESUMES_DIR

    master = load_master()

    result = {
        "ats_score":       0,
        "matched":         [],
        "missing":         [],
        "total":           0,
        "file":            None,
        "filename":        None,
        "below_threshold": False,
        "error":           None,
    }

    try:
        content = tailor_docx_content(jd_text, master, company, job_title, api_key)
    except Exception as exc:
        result["error"] = f"AI tailoring failed: {exc}. Generating with master resume data."
        content = {
            "summary": (
                f"Results-oriented professional seeking a {job_title} role at {company}. "
                "Brings proven experience in customer service, technical support, and operational "
                "workflows with a strong track record of accuracy and cross-functional collaboration."
            ),
            "competencies": master.get("skills", []),
            "bullets_by_role": {},
            "include_projects": False,
        }

    # Score against the tailored content, not the raw master resume
    tailored_blob = " ".join([
        content.get("summary", ""),
        " ".join(content.get("competencies", [])),
        " ".join(
            b for bullets in content.get("bullets_by_role", {}).values()
            for b in bullets
        ),
    ])
    ats_data = ats_score(jd_text, master, tailored_text=tailored_blob, company=company)
    result["ats_score"]       = ats_data["score"]
    result["matched"]         = ats_data["matched"]
    result["missing"]         = ats_data["missing"]
    result["total"]           = ats_data["total"]
    result["below_threshold"] = ats_data["score"] < 80

    try:
        out_path = build_docx(content, master, company, job_title, ats_data, out_dir)
        result["file"]     = out_path
        result["filename"] = out_path.name
    except Exception as exc:
        result["error"] = (result.get("error") or "") + f" DOCX build failed: {exc}"

    return result
